"""
Builds and sends the daily "what did the email poller learn today" report —
a .docx summarizing every fact email_learning.py extracted from real client
emails and staff replies that day, emailed to the studio owner at 18:00
Asia/Jerusalem. See studio_operations_and_communication_notes.md §11 and
app.py's _run_daily_report_scheduler().

A real, immediate send (not a draft) — this is an internal report to the
studio owner himself, not a client-facing message, same category as the
existing "flag as rush" internal notification in email_sender.py.
"""
import json
import tempfile
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from docx.shared import Pt

import email_learning
import email_sender

REPORT_TIMEZONE = ZoneInfo("Asia/Jerusalem")


def _entries_for_date(local_date):
    """local_date: a date object (in REPORT_TIMEZONE). Returns log entries
    (dicts with fact/citation/logged_at) whose logged_at falls on that
    local date."""
    if not email_learning.LOG_PATH.exists():
        return []
    entries = []
    with email_learning.LOG_PATH.open(encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                row = json.loads(line)
            except json.JSONDecodeError:
                continue
            try:
                from datetime import datetime
                logged_at = datetime.fromisoformat(row["logged_at"])
            except (KeyError, ValueError):
                continue
            if logged_at.astimezone(REPORT_TIMEZONE).date() == local_date:
                entries.append(row)
    return entries


def build_report_docx(local_date, entries):
    """Writes a .docx to a temp file and returns its path."""
    doc = Document()
    title = doc.add_heading(f"Daily Learning Summary — {local_date.isoformat()}", level=1)

    intro = doc.add_paragraph()
    intro.add_run(
        "New, durable facts the email poller extracted today from real client emails "
        "and the studio's actual replies (see studio_operations_and_communication_notes.md "
        "§10 — these are pending review, not yet folded into the main notes)."
    ).italic = True

    if not entries:
        p = doc.add_paragraph("Nothing new learned today.")
        p.runs[0].font.size = Pt(11)
    else:
        for entry in entries:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(entry["fact"])
            run.font.size = Pt(11)
            citation_p = doc.add_paragraph()
            citation_run = citation_p.add_run(f"  ({entry.get('citation', '')})")
            citation_run.italic = True
            citation_run.font.size = Pt(9)

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    return tmp.name


def send_daily_report(local_date, recipient):
    entries = _entries_for_date(local_date)
    path = build_report_docx(local_date, entries)
    try:
        subject = f"Daily learning summary — {local_date.isoformat()} ({len(entries)} new)"
        body = (
            f"Attached: today's summary of new things learned from real client emails and "
            f"replies ({len(entries)} new fact(s)).\n\n"
            "Sent automatically by the email poller."
        )
        filename = f"learning-summary-{local_date.isoformat()}.docx"
        email_sender.send_email(recipient, subject, body, attachments=[(path, filename)])
    finally:
        Path(path).unlink(missing_ok=True)
